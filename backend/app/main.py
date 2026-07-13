import asyncio
import csv
import hashlib
import io
import json
import re
import uuid
import zipfile
from collections import Counter
from datetime import datetime, timezone
from pathlib import Path, PurePosixPath
from typing import Any

from fastapi import BackgroundTasks, FastAPI, File, Form, HTTPException, UploadFile
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import StreamingResponse

from .anonymizer import redact_text
from .config import BASE_DIR, settings
from .instruction_parser import parse_instruction
from .knowledge_graph import knowledge_graph
from .pipeline import run_pipeline
from .semantic_adapter import semantic_encoder
from .schemas import (
    DetectRequest, DetectResponse, EntityType, FinalTextUpdate, InstructionRequest, KnowledgeLookupRequest,
    PolicyUpdate, ProcessingConfig, ProjectCreate, ProjectUpdate, RedactRequest,
    ReviewRequest, RuleCreate, RuleUpdate, Strategy,
)
from .storage import RevisionConflictError, Storage

app = FastAPI(title=settings.app_name, version="0.2.0", docs_url="/api/docs")
app.add_middleware(CORSMiddleware, allow_origins=settings.origins, allow_origin_regex=r"^(chrome-extension|moz-extension)://.*$", allow_credentials=True, allow_methods=["*"], allow_headers=["*"])
storage = Storage(settings.database_path)

DEFAULT_POLICIES = {key.value: "mask" for key in EntityType}
DEMO_METRICS = {
    "is_demo": True,
    "notice": "演示数据：接入冻结测试集后由 benchmark 输出自动替换",
    "metadata": {"source": "built-in-placeholder", "verified": False, "metric": "illustrative-only"},
    "systems": [
        {"name": "仅规则", "precision": 0.98, "recall": 0.61, "f1": 0.75, "latency": 18},
        {"name": "轻量 NER", "precision": 0.82, "recall": 0.76, "f1": 0.79, "latency": 142},
        {"name": "规则 + NER", "precision": 0.91, "recall": 0.87, "f1": 0.89, "latency": 168},
        {"name": "级联 + 14B", "precision": 0.93, "recall": 0.94, "f1": 0.935, "latency": 1840},
    ],
    "categories": [
        {"name": "电话", "recall": 0.99}, {"name": "邮箱", "recall": 0.99}, {"name": "姓名", "recall": 0.91},
        {"name": "地址", "recall": 0.89}, {"name": "机构", "recall": 0.92}, {"name": "证件", "recall": 0.98},
    ],
}


def _risk(spans) -> tuple[int, str]:
    active = [span for span in spans if span.status != "rejected"]
    score = min(100, sum(18 if span.entity_type in {EntityType.ID_CARD, EntityType.BANK_CARD, EntityType.PASSPORT} else 10 for span in active))
    return score, "high" if score >= 60 else "medium" if score >= 25 else "low"


def _parse_text_records(filename: str, content: bytes) -> list[dict[str, Any]]:
    suffix = Path(filename).suffix.lower()
    if suffix in {".txt", ".md", ".csv", ".json"}:
        try:
            decoded = content.decode("utf-8-sig")
        except UnicodeDecodeError as exc:
            raise HTTPException(400, f"{filename} 需使用 UTF-8 编码") from exc
        if suffix == ".csv":
            rows = list(csv.DictReader(io.StringIO(decoded)))
            if not rows:
                raise HTTPException(400, f"{filename} 为空")
            column = "text" if "text" in rows[0] else "content" if "content" in rows[0] else next(iter(rows[0]))
            return [{"file": filename, "row": index + 1, "text": row.get(column, "")} for index, row in enumerate(rows) if row.get(column)]
        if suffix == ".json":
            try:
                parsed = json.loads(decoded)
            except json.JSONDecodeError as exc:
                raise HTTPException(400, f"JSON 格式错误（{filename} 第 {exc.lineno} 行）") from exc
            if not isinstance(parsed, list):
                raise HTTPException(400, f"{filename} 的 JSON 顶层需为数组")
            records = []
            for index, item in enumerate(parsed):
                value = str(item.get("text", item.get("content", ""))) if isinstance(item, dict) else str(item)
                if value.strip():
                    records.append({"file": filename, "row": index + 1, "text": value})
            return records
        return [{"file": filename, "row": index + 1, "text": line} for index, line in enumerate(decoded.splitlines()) if line.strip()]
    if suffix == ".docx":
        try:
            from docx import Document
        except ImportError as exc:
            raise HTTPException(503, "服务器尚未安装 python-docx，无法读取 DOCX") from exc
        document = Document(io.BytesIO(content))
        parts = [paragraph.text for paragraph in document.paragraphs if paragraph.text.strip()]
        for table in document.tables:
            parts.extend("\t".join(cell.text for cell in row.cells) for row in table.rows)
        text = "\n".join(item for item in parts if item.strip())
        return [{"file": filename, "row": 1, "text": text}] if text else []
    if suffix == ".pdf":
        try:
            from pypdf import PdfReader
        except ImportError as exc:
            raise HTTPException(503, "服务器尚未安装 pypdf，无法读取 PDF") from exc
        reader = PdfReader(io.BytesIO(content))
        return [{"file": filename, "row": index + 1, "text": text} for index, page in enumerate(reader.pages) if (text := (page.extract_text() or "").strip())]
    raise HTTPException(415, f"暂不支持 {suffix or '无扩展名'} 文件：{filename}")


async def _process_batch(job_id: str, records: list[dict[str, Any]], config: ProcessingConfig, project_id: str | None = None):
    existing = storage.get_job(job_id) or {}
    payload = {**existing.get("payload", {}), "results": [], "failures": [], "config": config.model_dump(mode="json")}
    persistent_rules = await asyncio.to_thread(storage.list_rules, project_id)
    policies = None
    if config.use_policies:
        raw_policies = await asyncio.to_thread(storage.get_setting, "policies", DEFAULT_POLICIES)
        policies = {EntityType(key): Strategy(value) for key, value in raw_policies.items()}
    storage.update_job(job_id, status="running", payload=payload)
    failed = 0
    for index, record in enumerate(records):
        try:
            request = DetectRequest(text=record["text"], **config.model_dump(mode="json"))
            spans, redacted, trace, applied_config = await run_pipeline(request, policies, persistent_rules)
            task_id = f"task_{uuid.uuid4().hex[:12]}"
            active_spans = [span for span in spans if span.status != "rejected"]
            counts = Counter(span.entity_type.value for span in active_spans)
            _, risk = _risk(spans)
            created = datetime.now(timezone.utc).isoformat()
            snapshot = DetectResponse(
                task_id=task_id, text=record["text"], spans=spans, redacted_text=redacted, trace=trace,
                summary={"total": len(active_spans), "pending": sum(span.status == "pending" for span in active_spans), "risk_score": _risk(spans)[0], "by_type": counts},
                model={"name": settings.llm_model, "enabled": settings.llm_enabled, "mode": "non-thinking", "runtime": "vLLM" if settings.llm_enabled else "lightweight"},
                created_at=created, final_text=redacted, final_revision=0, has_manual_edits=False,
                applied_config=applied_config, project_id=project_id,
            )
            await asyncio.to_thread(storage.save_task, task_id, redacted[:80], len(active_spans), risk, snapshot.model_dump(mode="json"))
            payload["results"].append({
                "file": record["file"], "row": record["row"], "task_id": task_id,
                "redacted_text": redacted, "final_text": redacted, "final_revision": 0,
                "entity_count": len(active_spans),
                "pending_count": sum(span.status == "pending" for span in active_spans),
                "status": "needs_review" if any(span.status == "pending" or span.conflict for span in spans) else "completed",
                "applied_config": applied_config,
            })
        except Exception as exc:
            failed += 1
            raw_text = record["text"]
            payload["failures"].append({
                "file": record["file"], "row": record["row"],
                "text_length": len(raw_text),
                "text_hash": hashlib.sha256(raw_text.encode("utf-8")).hexdigest(),
                "error": f"{type(exc).__name__}: {exc}",
            })
        storage.update_job(job_id, status="running", processed=index + 1, failed=failed, payload=payload)
    storage.update_job(job_id, status="completed_with_errors" if failed else "completed", processed=len(records), failed=failed, payload=payload)


def _safe_export_name(filename: str) -> str:
    parts = [part for part in PurePosixPath(filename.replace("\\", "/")).parts if part not in {"", ".", ".."}]
    safe = PurePosixPath(*parts) if parts else PurePosixPath("result.txt")
    stem = safe.stem or "result"
    return str(safe.with_name(f"{stem}.redacted.txt"))


def _public_job(job: dict[str, Any]) -> dict[str, Any]:
    """Return a UI/export-safe job view and hydrate each row from its latest task revision."""
    public = {**job, "payload": {**job.get("payload", {})}}
    results = []
    for item in job.get("payload", {}).get("results", []):
        clean = {key: value for key, value in item.items() if key != "text"}
        task = storage.get_task(str(item.get("task_id", ""))) if item.get("task_id") else None
        if task:
            clean["redacted_text"] = task.get("redacted_text", clean.get("redacted_text", ""))
            clean["final_text"] = task.get("final_text", clean["redacted_text"])
            clean["final_revision"] = int(task.get("final_revision", 0))
            clean["has_manual_edits"] = bool(task.get("has_manual_edits", False))
        else:
            clean["final_text"] = clean.get("final_text", clean.get("redacted_text", ""))
            clean["final_revision"] = int(clean.get("final_revision", 0))
            clean["has_manual_edits"] = clean["final_text"] != clean.get("redacted_text", "")
        results.append(clean)
    failures = []
    for item in job.get("payload", {}).get("failures", []):
        clean = {key: value for key, value in item.items() if key != "text"}
        raw_text = item.get("text")
        if raw_text is not None:
            clean.setdefault("text_length", len(raw_text))
            clean.setdefault("text_hash", hashlib.sha256(raw_text.encode("utf-8")).hexdigest())
        failures.append(clean)
    public["payload"]["results"] = results
    public["payload"]["failures"] = failures
    return public


@app.get("/api/v1/health")
def health():
    try:
        database = "online" if storage.ping() else "offline"
    except Exception:
        database = "offline"
    return {"status": "ok" if database == "online" else "degraded", "version": app.version, "mode": "llm" if settings.llm_enabled else "lightweight", "database": database, "deployment": "self-hosted"}


@app.get("/api/v1/models")
def models():
    return {
        "active": settings.llm_model, "enabled": settings.llm_enabled,
        "mode": "non-thinking", "provider": "OpenAI-compatible vLLM" if settings.llm_enabled else "离线轻量模式",
        "candidates": ["Qwen/Qwen3-14B", "Qwen/Qwen3-14B-AWQ", "Qwen/Qwen3-8B", "Qwen/Qwen2.5-7B-Instruct"],
        "ner": settings.ner_model, "multilingual": True,
        "semantic": semantic_encoder.status(),
        "knowledge_graph": knowledge_graph.status(),
    }


@app.get("/api/v1/knowledge/status")
def knowledge_status():
    return knowledge_graph.status()


@app.post("/api/v1/knowledge/lookup")
async def knowledge_lookup(request: KnowledgeLookupRequest):
    return (await knowledge_graph.lookup(request.term, request.entity_type, request.allow_remote)).model_dump()


@app.post("/api/v1/instructions/parse")
async def parse_requirement(request: InstructionRequest):
    return await parse_instruction(request.instruction, request.use_llm)


@app.post("/api/v1/detect", response_model=DetectResponse)
async def detect(request: DetectRequest):
    if request.project_id:
        project = await asyncio.to_thread(storage.get_project, request.project_id)
        if project is None:
            raise HTTPException(404, "项目不存在")
        explicit = request.model_dump(include=request.model_fields_set, mode="json")
        request = DetectRequest.model_validate({**project["config"], **explicit})
    task_id = f"task_{uuid.uuid4().hex[:12]}"
    policies = None
    if request.use_policies:
        raw_policies = await asyncio.to_thread(storage.get_setting, "policies", DEFAULT_POLICIES)
        policies = {EntityType(key): Strategy(value) for key, value in raw_policies.items()}
    persistent_rules = await asyncio.to_thread(storage.list_rules, request.project_id)
    spans, redacted, trace, applied_config = await run_pipeline(request, policies, persistent_rules)
    active_spans = [span for span in spans if span.status != "rejected"]
    counts = Counter(span.entity_type.value for span in active_spans)
    risk_score, risk = _risk(spans)
    created = datetime.now(timezone.utc).isoformat()
    response = DetectResponse(
        task_id=task_id, text=request.text, spans=spans, redacted_text=redacted, trace=trace,
        summary={"total": len(active_spans), "pending": sum(span.status == "pending" for span in active_spans), "risk_score": risk_score, "by_type": counts},
        model={"name": settings.llm_model, "enabled": settings.llm_enabled, "mode": "non-thinking", "runtime": "vLLM" if settings.llm_enabled else "lightweight"},
        created_at=created, final_text=redacted, final_revision=0, has_manual_edits=False, applied_config=applied_config, project_id=request.project_id,
    )
    await asyncio.to_thread(storage.save_task, task_id, redacted[:80], len(active_spans), risk, response.model_dump(mode="json"))
    return response


@app.post("/api/v1/redact")
def redact(request: RedactRequest):
    for span in request.spans:
        if span.end > len(request.text) or request.text[span.start:span.end] != span.text:
            raise HTTPException(422, f"Span {span.id} 与原文偏移不一致")
    allowed_statuses = {"accepted", "pending"} if request.risk_level == "strict" else {"accepted"}
    active = sorted((span for span in request.spans if span.status in allowed_statuses), key=lambda span: (span.start, span.end))
    if any(current.start < previous.end for previous, current in zip(active, active[1:])):
        raise HTTPException(422, "有效 Span 之间不能重叠，请先完成冲突消歧")
    return {"redacted_text": redact_text(request.text, request.spans, request.strategy, request.privacy_strength, include_pending=request.risk_level == "strict")}


@app.post("/api/v1/reviews")
async def review(request: ReviewRequest):
    try:
        snapshot = await asyncio.to_thread(storage.apply_review, request)
    except LookupError as exc:
        raise HTTPException(404, str(exc)) from exc
    except ValueError as exc:
        raise HTTPException(422, str(exc)) from exc
    if snapshot is None:
        raise HTTPException(404, "任务不存在")
    return {"ok": True, "recorded_at": datetime.now(timezone.utc).isoformat(), "snapshot": snapshot}


@app.get("/api/v1/reviews")
async def review_queue():
    return {"items": await asyncio.to_thread(storage.review_queue)}


@app.put("/api/v1/tasks/{task_id}/final-text")
async def save_final_text(task_id: str, request: FinalTextUpdate):
    try:
        result = await asyncio.to_thread(storage.save_final_text, task_id, request)
    except RevisionConflictError as exc:
        raise HTTPException(409, str(exc)) from exc
    if result is None:
        raise HTTPException(404, "任务不存在")
    return result


@app.get("/api/v1/tasks/{task_id}")
async def get_task(task_id: str):
    task = await asyncio.to_thread(storage.get_task, task_id)
    if task is None:
        raise HTTPException(404, "任务不存在")
    return task


@app.delete("/api/v1/tasks/{task_id}")
async def delete_task(task_id: str):
    if not await asyncio.to_thread(storage.delete_task, task_id):
        raise HTTPException(404, "任务不存在")
    return {"ok": True}


@app.delete("/api/v1/tasks")
async def purge_tasks(older_than_days: int = 30):
    if not 1 <= older_than_days <= 3650:
        raise HTTPException(422, "保留天数需在 1 到 3650 之间")
    return {"deleted": await asyncio.to_thread(storage.purge_tasks, older_than_days)}


@app.get("/api/v1/history")
async def history():
    items, audits = await asyncio.gather(asyncio.to_thread(storage.history), asyncio.to_thread(storage.audits))
    return {"items": items, "audits": audits}


@app.get("/api/v1/evaluations")
def evaluations():
    results_file = BASE_DIR / "reports" / "experiment_results" / "latest.json"
    if results_file.exists():
        result = json.loads(results_file.read_text(encoding="utf-8"))
        result["is_demo"] = False
        result.setdefault("metadata", {}).update(source=str(results_file), verified=True)
        return result
    return DEMO_METRICS


@app.get("/api/v1/policies")
async def get_policies():
    return {"policies": await asyncio.to_thread(storage.get_setting, "policies", DEFAULT_POLICIES), "version": "policy-2026.07"}


@app.put("/api/v1/policies")
async def update_policies(request: PolicyUpdate):
    policies = await asyncio.to_thread(storage.get_setting, "policies", DEFAULT_POLICIES.copy())
    policies.update({key.value: value.value for key, value in request.policies.items()})
    await asyncio.to_thread(storage.set_setting, "policies", policies)
    return {"policies": policies, "version": "policy-2026.07-custom"}


@app.get("/api/v1/projects")
async def list_projects():
    return {"items": await asyncio.to_thread(storage.list_projects)}


@app.post("/api/v1/projects")
async def create_project(request: ProjectCreate):
    project_id = f"project_{uuid.uuid4().hex[:10]}"
    return await asyncio.to_thread(storage.create_project, project_id, request)


@app.put("/api/v1/projects/{project_id}")
async def update_project(project_id: str, request: ProjectUpdate):
    result = await asyncio.to_thread(storage.update_project, project_id, request)
    if result is None:
        raise HTTPException(404, "项目不存在")
    return result


@app.delete("/api/v1/projects/{project_id}")
async def delete_project(project_id: str):
    try:
        deleted = await asyncio.to_thread(storage.delete_project, project_id)
    except ValueError as exc:
        raise HTTPException(409, str(exc)) from exc
    if not deleted:
        raise HTTPException(404, "项目不存在")
    return {"ok": True}


@app.post("/api/v1/extract")
async def extract_files(files: list[UploadFile] = File(...)):
    records: list[dict[str, Any]] = []
    for upload in files:
        content = await upload.read(settings.max_upload_bytes + 1)
        if len(content) > settings.max_upload_bytes:
            raise HTTPException(413, f"{upload.filename} 超过 {settings.max_upload_bytes // 1_000_000} MB")
        records.extend(_parse_text_records(upload.filename or "upload.txt", content))
    return {"records": records, "text": "\n".join(record["text"] for record in records), "files": len(files)}


@app.get("/api/v1/rules")
async def list_rules(project_id: str | None = None):
    return {"items": await asyncio.to_thread(storage.list_rules, project_id)}


@app.post("/api/v1/rules")
async def create_rule(request: RuleCreate):
    if request.project_id and await asyncio.to_thread(storage.get_project, request.project_id) is None:
        raise HTTPException(404, "项目不存在")
    if request.kind == "regex":
        try:
            re.compile(request.pattern)
        except re.error as exc:
            raise HTTPException(422, f"正则表达式无效：{exc}") from exc
    rule_id = f"rule_{uuid.uuid4().hex[:10]}"
    return await asyncio.to_thread(storage.save_rule, rule_id, request)

@app.put("/api/v1/rules/{rule_id}")
async def update_rule(rule_id: str, request: RuleUpdate):
    existing = await asyncio.to_thread(storage.get_rule, rule_id)
    if existing is None:
        raise HTTPException(404, "规则不存在")
    next_kind = request.kind or existing["kind"]
    next_pattern = request.pattern or existing["pattern"]
    if next_kind == "regex":
        try:
            re.compile(next_pattern)
        except re.error as exc:
            raise HTTPException(422, f"正则表达式无效：{exc}") from exc
    return await asyncio.to_thread(storage.update_rule, rule_id, request)


@app.delete("/api/v1/rules/{rule_id}")
async def delete_rule(rule_id: str):
    if not await asyncio.to_thread(storage.delete_rule, rule_id):
        raise HTTPException(404, "规则不存在")
    return {"ok": True}


@app.post("/api/v1/jobs")
async def create_batch_job(
    background_tasks: BackgroundTasks,
    files: list[UploadFile] | None = File(default=None),
    file: UploadFile | None = File(default=None),
    config_json: str | None = Form(default=None),
    project_id: str | None = Form(default=None),
    strategy: Strategy = Strategy.MASK,
):
    uploads = list(files or []) + ([file] if file else [])
    if not uploads:
        raise HTTPException(422, "请至少上传一个文件")
    project = None
    if project_id:
        project = await asyncio.to_thread(storage.get_project, project_id)
        if project is None:
            raise HTTPException(404, "项目不存在")
    records: list[dict[str, Any]] = []
    for upload in uploads:
        content = await upload.read(settings.max_upload_bytes + 1)
        if len(content) > settings.max_upload_bytes:
            raise HTTPException(413, f"{upload.filename} 超过 {settings.max_upload_bytes // 1_000_000} MB")
        records.extend(_parse_text_records(upload.filename or "upload.txt", content))
    if not records:
        raise HTTPException(400, "文件中没有可处理文本")
    if len(records) > settings.max_batch_records:
        raise HTTPException(413, f"单次最多处理 {settings.max_batch_records} 条记录")
    try:
        explicit_config = json.loads(config_json) if config_json else {"strategy": strategy.value}
        config_data = {**(project["config"] if project else {}), **explicit_config, "project_id": project_id}
        config = ProcessingConfig.model_validate(config_data)
    except (json.JSONDecodeError, ValueError) as exc:
        raise HTTPException(422, f"处理配置无效：{exc}") from exc
    for index, record in enumerate(records):
        if len(record["text"]) > 100_000:
            raise HTTPException(422, f"第 {index + 1} 条文本超过 100000 字符")
    job_id = f"job_{uuid.uuid4().hex[:10]}"
    await asyncio.to_thread(storage.create_job, job_id, project_id, len(records), {"results": [], "failures": [], "files": [upload.filename for upload in uploads], "config": config.model_dump(mode="json")})
    background_tasks.add_task(_process_batch, job_id, records, config, project_id)
    return await asyncio.to_thread(storage.get_job, job_id)


@app.get("/api/v1/jobs")
async def list_jobs():
    jobs = await asyncio.to_thread(storage.list_jobs)
    return {"items": await asyncio.gather(*(asyncio.to_thread(_public_job, job) for job in jobs))}


@app.get("/api/v1/jobs/{job_id}/download")
async def download_job(job_id: str):
    stored_job = await asyncio.to_thread(storage.get_job, job_id)
    if stored_job is None:
        raise HTTPException(404, "批处理任务不存在")
    job = await asyncio.to_thread(_public_job, stored_job)
    if job["status"] in {"queued", "running"}:
        raise HTTPException(409, "批处理尚未完成")
    buffer = io.BytesIO()
    grouped: dict[str, list[dict[str, Any]]] = {}
    for item in job["payload"].get("results", []):
        grouped.setdefault(item.get("file") or "result.txt", []).append(item)
    with zipfile.ZipFile(buffer, "w", compression=zipfile.ZIP_DEFLATED) as archive:
        used_names: set[str] = set()
        for filename, items in grouped.items():
            export_name = _safe_export_name(filename)
            candidate = export_name
            suffix = 2
            while candidate in used_names:
                path = PurePosixPath(export_name)
                candidate = str(path.with_name(f"{path.stem}-{suffix}{path.suffix}"))
                suffix += 1
            used_names.add(candidate)
            ordered = sorted(items, key=lambda item: int(item.get("row", 0)))
            archive.writestr(candidate, "\n".join(str(item.get("final_text", item.get("redacted_text", ""))) for item in ordered))
        archive.writestr("manifest.json", json.dumps(job, ensure_ascii=False, indent=2))
        failures = job["payload"].get("failures", [])
        if failures:
            output = io.StringIO()
            writer = csv.DictWriter(output, fieldnames=["file", "row", "error", "text_length", "text_hash"], extrasaction="ignore")
            writer.writeheader()
            writer.writerows(failures)
            archive.writestr("failures.csv", "\ufeff" + output.getvalue())
    buffer.seek(0)
    headers = {"Content-Disposition": f'attachment; filename="{job_id}-redacted.zip"'}
    return StreamingResponse(buffer, media_type="application/zip", headers=headers)


@app.get("/api/v1/jobs/{job_id}")
async def get_job(job_id: str):
    result = await asyncio.to_thread(storage.get_job, job_id)
    if result is None:
        raise HTTPException(404, "批处理任务不存在")
    return await asyncio.to_thread(_public_job, result)


@app.delete("/api/v1/jobs/{job_id}")
async def delete_job(job_id: str):
    if not await asyncio.to_thread(storage.delete_job, job_id):
        raise HTTPException(404, "批处理任务不存在")
    return {"ok": True}
